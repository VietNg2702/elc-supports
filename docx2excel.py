"""Convert DOCX content (text and tables) to an Excel worksheet.

Usage:
	python pdf2excel.py --docx "document.docx" --output "document.xlsx"

Dependencies:
	pip install openpyxl python-docx
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.oxml.ns import qn


def normalize_hex_color(color: str | None) -> str:
	"""Normalize DOCX color values to Excel-compatible RGB hex."""
	if not color or color.lower() == "auto":
		return "000000"
	color = color.strip().lstrip("#").upper()
	if len(color) == 3:
		return "".join(char * 2 for char in color)
	if len(color) == 6:
		return color
	return "000000"


def find_border_element(parent: Any, side_tag: str) -> Any:
	"""Find a border XML node by side name if it exists."""
	if parent is None:
		return None
	for child in parent.iterchildren():
		if child.tag == qn(side_tag):
			return child
	return None


def read_border_definitions(table: Any, cell: Any, side: str) -> tuple[bool, str]:
	"""Read border visibility and color from cell or table-level DOCX borders."""
	side_map = {
		"left": "w:left",
		"right": "w:right",
		"top": "w:top",
		"bottom": "w:bottom",
	}
	side_tag = side_map[side]

	tc_pr = cell._tc.tcPr
	tc_borders = tc_pr.first_child_found_in("w:tcBorders") if tc_pr is not None else None
	tbl_pr = table._tbl.tblPr
	tbl_borders = tbl_pr.first_child_found_in("w:tblBorders") if tbl_pr is not None else None

	node = find_border_element(tc_borders, side_tag)
	if node is None:
		node = find_border_element(tbl_borders, side_tag)

	if node is None:
		return False, "000000"

	val = node.get(qn("w:val"), "single").lower()
	if val in {"nil", "none"}:
		return False, "000000"

	color = normalize_hex_color(node.get(qn("w:color")))
	return True, color


def build_excel_border(table: Any, cell: Any) -> Border:
	"""Build an Excel border from DOCX border metadata."""
	left_on, left_color = read_border_definitions(table, cell, "left")
	right_on, right_color = read_border_definitions(table, cell, "right")
	top_on, top_color = read_border_definitions(table, cell, "top")
	bottom_on, bottom_color = read_border_definitions(table, cell, "bottom")

	return Border(
		left=Side(style="thin", color=left_color) if left_on else Side(style=None),
		right=Side(style="thin", color=right_color) if right_on else Side(style=None),
		top=Side(style="thin", color=top_color) if top_on else Side(style=None),
		bottom=Side(style="thin", color=bottom_color) if bottom_on else Side(style=None),
	)


def write_docx_tables_to_excel(docx_path: Path, output_path: Path) -> None:
	"""Write all DOCX content (text + tables) into one Excel sheet."""
	doc = Document(str(docx_path))
	workbook = Workbook()
	sheet = workbook.active
	sheet.title = "Content"

	row_offset = 1
	max_table_columns = max((len(row.cells) for table in doc.tables for row in table.rows), default=1)
	content_span_columns = max(4, max_table_columns)
	max_column = content_span_columns

	table_lookup = {id(table._tbl): table for table in doc.tables}
	paragraph_lookup = {id(paragraph._p): paragraph for paragraph in doc.paragraphs}
	table_index = 0
	added_content = False

	for element in doc.element.body.iterchildren():
		if element.tag == qn("w:p"):
			paragraph = paragraph_lookup.get(id(element))
			if paragraph is None:
				continue

			text = "\n".join(part.strip() for part in paragraph.text.splitlines() if part.strip())
			if not text:
				continue

			sheet.cell(row=row_offset, column=1, value=text)
			sheet.cell(row=row_offset, column=1).alignment = Alignment(wrap_text=True, vertical="top")
			if content_span_columns > 1:
				sheet.merge_cells(
					start_row=row_offset,
					start_column=1,
					end_row=row_offset,
					end_column=content_span_columns,
				)

			sheet.row_dimensions[row_offset].height = 24
			row_offset += 1
			added_content = True
			continue

		if element.tag != qn("w:tbl"):
			continue

		table = table_lookup.get(id(element))
		if table is None:
			continue

		table_index += 1
		sheet.cell(row=row_offset, column=1, value=f"Table {table_index}")
		sheet.cell(row=row_offset, column=1).alignment = Alignment(vertical="top")
		if content_span_columns > 1:
			sheet.merge_cells(
				start_row=row_offset,
				start_column=1,
				end_row=row_offset,
				end_column=content_span_columns,
			)
		row_offset += 1

		cell_refs: list[list[Any]] = []
		cell_ids: list[list[int]] = []
		for row in table.rows:
			row_cells = list(row.cells)
			cell_refs.append(row_cells)
			cell_ids.append([id(cell._tc) for cell in row_cells])

		first_positions: dict[int, tuple[int, int]] = {}
		span_boxes: dict[int, tuple[int, int, int, int]] = {}

		for r, row_cells in enumerate(cell_refs):
			for c, cell in enumerate(row_cells):
				cid = id(cell._tc)
				if cid not in first_positions:
					first_positions[cid] = (r, c)
				if cid not in span_boxes:
					span_boxes[cid] = (r, c, r, c)
				else:
					r0, c0, r1, c1 = span_boxes[cid]
					span_boxes[cid] = (min(r0, r), min(c0, c), max(r1, r), max(c1, c))

		for r, row_cells in enumerate(cell_refs):
			for c, cell in enumerate(row_cells):
				cid = id(cell._tc)
				if first_positions[cid] != (r, c):
					continue

				excel_row = row_offset + r
				excel_col = c + 1
				value = "\n".join(part.strip() for part in cell.text.splitlines() if part.strip())
				target = sheet.cell(row=excel_row, column=excel_col, value=value)
				target.alignment = Alignment(wrap_text=True, vertical="top")
				target.border = build_excel_border(table, cell)

		for cid, (r0, c0, r1, c1) in span_boxes.items():
			if r1 == r0 and c1 == c0:
				continue

			full_rect = True
			for rr in range(r0, r1 + 1):
				for cc in range(c0, c1 + 1):
					if cell_ids[rr][cc] != cid:
						full_rect = False
						break
				if not full_rect:
					break

			if not full_rect:
				continue

			sheet.merge_cells(
				start_row=row_offset + r0,
				start_column=c0 + 1,
				end_row=row_offset + r1,
				end_column=c1 + 1,
			)

		table_rows = len(cell_refs)
		table_cols = max((len(row_cells) for row_cells in cell_refs), default=1)
		for r in range(row_offset, row_offset + table_rows):
			sheet.row_dimensions[r].height = 22

		row_offset += table_rows + 2
		max_column = max(max_column, table_cols)
		added_content = True

	if not added_content:
		sheet.cell(row=1, column=1, value="No readable content found in DOCX.")

	for col in range(1, max_column + 1):
		sheet.column_dimensions[get_column_letter(col)].width = 22

	workbook.save(output_path)


def write_to_excel(input_docx: Path, output_path: Path) -> None:
	"""Convert DOCX to a one-sheet Excel file."""
	write_docx_tables_to_excel(input_docx, output_path)


def build_parser() -> argparse.ArgumentParser:
	parser = argparse.ArgumentParser(
		description="Convert DOCX content to Excel in one worksheet."
	)
	parser.add_argument(
		"--docx",
		required=True,
		type=Path,
		help="Path to DOCX file for extraction.",
	)
	parser.add_argument(
		"--output",
		type=Path,
		help="Path to the output XLSX file.",
	)
	return parser


def main() -> None:
	parser = build_parser()
	args = parser.parse_args()

	input_docx: Path = args.docx

	if not input_docx.exists() or input_docx.suffix.lower() != ".docx":
		raise ValueError(f"Invalid DOCX path: {input_docx}")

	if args.output is not None:
		output_path = args.output
	else:
		output_path = input_docx.with_suffix(".xlsx")

	write_to_excel(input_docx, output_path)

	print(f"Converted to '{output_path}'")


if __name__ == "__main__":
	main()
